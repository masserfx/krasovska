#!/usr/bin/env python3
"""
Vytvoří reference.docx šablonu pro pandoc s vlastními fonty a barvami,
pak exportuje CEO_BRIEFING.md jako finální .docx.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess, os

DOCS_DIR = "/Users/lhradek/code/work/krasovska/questionnaire/docs"
TEMPLATE  = f"{DOCS_DIR}/word_template.docx"
OUTPUT    = f"{DOCS_DIR}/CEO_BRIEFING.docx"
SOURCE_MD = f"{DOCS_DIR}/CEO_BRIEFING.md"

# ── Barvy ──────────────────────────────────────────────────────────────────
H1_COLOR   = RGBColor(0x1B, 0x3A, 0x5C)   # tmavě modrá
H2_COLOR   = RGBColor(0x1B, 0x5C, 0x38)   # tmavě zelená
H3_COLOR   = RGBColor(0x4A, 0x4A, 0x4A)   # tmavě šedá
H4_COLOR   = RGBColor(0x6B, 0x6B, 0x6B)
ACCENT     = RGBColor(0xD4, 0x45, 0x06)   # oranžovo-červená (sportovní akcent)
TABLE_HEAD = RGBColor(0x1B, 0x3A, 0x5C)   # hlavička tabulky = H1
TABLE_ROW  = RGBColor(0xEA, 0xF0, 0xF6)   # každý druhý řádek

# ── Font ───────────────────────────────────────────────────────────────────
BODY_FONT = "Calibri"
HEAD_FONT = "Calibri"

def set_run_font(run, name, size_pt, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def style_paragraph(style, font_name, size_pt, bold=False, space_before=0,
                     space_after=6, color=None, keep_with_next=False):
    f = style.font
    f.name = font_name
    f.size = Pt(size_pt)
    f.bold = bold
    if color:
        f.color.rgb = color
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.keep_with_next = keep_with_next

def shade_cell(cell, hex_color):
    """Vyplní buňku barvou pozadí."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

# ── Vytvoření šablony ──────────────────────────────────────────────────────
doc = Document()

# Okraje stránky
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# Výchozí styl odstavce (Normal)
normal = doc.styles["Normal"]
style_paragraph(normal, BODY_FONT, 10.5, space_after=4)

# Heading 1
h1 = doc.styles["Heading 1"]
style_paragraph(h1, HEAD_FONT, 18, bold=True, space_before=14, space_after=6,
                color=H1_COLOR, keep_with_next=True)
h1.paragraph_format.border_bottom = None   # pandoc přidá border jinak

# Heading 2
h2 = doc.styles["Heading 2"]
style_paragraph(h2, HEAD_FONT, 13, bold=True, space_before=12, space_after=4,
                color=H2_COLOR, keep_with_next=True)

# Heading 3
h3 = doc.styles["Heading 3"]
style_paragraph(h3, HEAD_FONT, 11, bold=True, space_before=10, space_after=3,
                color=H3_COLOR, keep_with_next=True)

# Heading 4
h4 = doc.styles["Heading 4"]
style_paragraph(h4, HEAD_FONT, 10.5, bold=True, space_before=8, space_after=2,
                color=H4_COLOR, keep_with_next=True)

# Block quote (pro > citáty)
if "Block Text" in doc.styles:
    bq = doc.styles["Block Text"]
    style_paragraph(bq, BODY_FONT, 10, italic=True, space_before=4, space_after=4)
    bq.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# Body Text
if "Body Text" in doc.styles:
    style_paragraph(doc.styles["Body Text"], BODY_FONT, 10.5, space_after=4)

# List styles
for lst_name in ("List Bullet", "List Bullet 2", "List Number"):
    if lst_name in doc.styles:
        style_paragraph(doc.styles[lst_name], BODY_FONT, 10.5, space_after=2)

doc.save(TEMPLATE)
print(f"Šablona uložena: {TEMPLATE}")

# ── Pandoc export ──────────────────────────────────────────────────────────
cmd = [
    "pandoc", SOURCE_MD,
    "-o", OUTPUT,
    "--from=markdown",
    "--to=docx",
    f"--reference-doc={TEMPLATE}",
    "--metadata", "title=CEO Briefing – Bistro Hala Krašovská",
    "--metadata", "lang=cs-CZ",
]
result = subprocess.run(cmd, capture_output=True, text=True)
if result.returncode != 0:
    print("CHYBA:", result.stderr)
else:
    print(f"Dokument uložen: {OUTPUT}")

# ── Post-process: nastylovat tabulky ──────────────────────────────────────
from docx import Document as D2
doc2 = D2(OUTPUT)

for table in doc2.tables:
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
                for run in para.runs:
                    run.font.name = BODY_FONT
                    run.font.size = Pt(9.5)
            if i == 0:
                # Hlavička – tmavě modrá, bílý text, tučně
                shade_cell(cell, "1B3A5C")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
            elif i % 2 == 0:
                shade_cell(cell, "EAF0F6")
            else:
                shade_cell(cell, "FFFFFF")

doc2.save(OUTPUT)
print("Tabulky nastylované. Hotovo.")
