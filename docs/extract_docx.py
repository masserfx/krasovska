#!/usr/bin/env python3
"""Extrahuje text z .docx a uloží jako markdown."""
from docx import Document
import re

doc = Document("/Users/lhradek/code/work/krasovska/questionnaire/docs/CEO_BRIEFING.docx")

lines = []
for para in doc.paragraphs:
    style = para.style.name
    text  = para.text.strip()
    if not text:
        lines.append("")
        continue
    if style == "Heading 1":
        lines.append(f"# {text}")
    elif style == "Heading 2":
        lines.append(f"## {text}")
    elif style == "Heading 3":
        lines.append(f"### {text}")
    elif style == "Heading 4":
        lines.append(f"#### {text}")
    elif style.startswith("List Bullet"):
        lines.append(f"- {text}")
    elif style.startswith("List Number"):
        lines.append(f"1. {text}")
    elif style in ("Block Text", "Quote", "Intense Quote"):
        lines.append(f"> {text}")
    else:
        lines.append(text)

# Tabulky
full_text = "\n".join(lines)
print(full_text[:3000])
print("\n\n--- TABULKY ---")
for i, table in enumerate(doc.tables):
    print(f"\nTabulka {i+1}:")
    for row in table.rows:
        cells = [c.text.strip().replace("\n"," ") for c in row.cells]
        print(" | ".join(cells))
